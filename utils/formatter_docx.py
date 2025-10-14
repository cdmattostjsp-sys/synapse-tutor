"""
utils/formatter_docx.py (v3.5)
-------------------------------
Conversor Markdown → DOCX + PDF institucional (Synapse Tutor v2)

Novidades:
- Geração de rascunho numerado (ex: "Rascunho nº 003/2025")
- Inclusão de campo 'Responsável pela elaboração'
- Exportação simultânea em PDF (via pypandoc)
- Manutenção de cabeçalho, rodapé e resumo de análise
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
import os
import pypandoc
from datetime import datetime


def get_next_rascunho_number(rasc_dir: str) -> int:
    """Calcula o próximo número de rascunho com base nos arquivos existentes."""
    existing = [f for f in os.listdir(rasc_dir) if f.startswith("DFD_")]
    return len(existing) + 1


def markdown_to_docx(
    markdown_text: str,
    title: str = "Documento Institucional",
    summary_text: str = "",
    rasc_dir: str = "./exports/rascunhos",
    responsavel: str = "Não informado",
) -> tuple[BytesIO, str]:
    """
    Converte markdown → DOCX institucional + gera PDF.
    Retorna buffer DOCX e caminho do PDF.
    """

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Cabeçalho institucional
    header = section.header.paragraphs[0]
    header.text = "Tribunal de Justiça de São Paulo – SAAB | Synapse Tutor v2"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.style.font.size = Pt(9)
    header.style.font.color.rgb = RGBColor(90, 90, 90)

    # Rodapé institucional
    footer = section.footer.paragraphs[0]
    footer.text = "Documento gerado automaticamente — Synapse.IA | SAAB | TJSP"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style.font.size = Pt(8)
    footer.style.font.color.rgb = RGBColor(100, 100, 100)

    # Numeração sequencial
    numero = get_next_rascunho_number(rasc_dir)
    ano = datetime.now().year
    doc.add_heading(f"{title} – Rascunho nº {numero:03d}/{ano}", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Responsável pela elaboração
    p_info = doc.add_paragraph()
    p_info.add_run(f"Responsável pela elaboração: {responsavel}\n").bold = True
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Limpeza de markdown básico
    markdown_text = (
        markdown_text.replace("<b>", "").replace("</b>", "")
        .replace("<i>", "").replace("</i>", "")
        .replace("**", "").replace("*", "")
    )

    lines = markdown_text.splitlines()
    for line in lines:
        p = None
        text = line.strip()
        if not text:
            doc.add_paragraph()
            continue

        if text.startswith("# "):
            doc.add_heading(text[2:].strip(), level=1)
        elif text.startswith("## "):
            doc.add_heading(text[3:].strip(), level=2)
        elif text.startswith("### "):
            doc.add_heading(text[4:].strip(), level=3)
        elif text.startswith("💡"):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        elif text.startswith("- "):
            p = doc.add_paragraph(text[2:].strip(), style="List Bullet")
        elif text.startswith("---"):
            sep = doc.add_paragraph("―" * 30)
            sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p = doc.add_paragraph(text)

        if p:
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25

    # Resumo da Análise
    if summary_text:
        doc.add_page_break()
        doc.add_heading("📊 Resumo da Análise", level=2)
        doc.add_paragraph(summary_text)

    # Exportações
    os.makedirs(rasc_dir, exist_ok=True)
    file_base = f"DFD_{datetime.now():%Y%m%d_%H%M%S}"
    docx_path = os.path.join(rasc_dir, f"{file_base}.docx")
    pdf_path = os.path.join(rasc_dir, f"{file_base}.pdf")

    buffer = BytesIO()
    doc.save(buffer)
    with open(docx_path, "wb") as f:
        f.write(buffer.getvalue())

    # Gera PDF a partir do markdown
    try:
        pypandoc.convert_text(markdown_text, "pdf", format="md", outputfile=pdf_path, extra_args=["--standalone"])
    except Exception:
        pass  # não quebra fluxo se pypandoc não estiver instalado

    buffer.seek(0)
    return buffer, pdf_path
